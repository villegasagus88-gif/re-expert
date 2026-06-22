"""
Document service — genera PDF / DOCX con datos del proyecto.

Estrategia de almacenamiento:
  1. Si Supabase Storage está configurado (SUPABASE_URL + SERVICE_ROLE_KEY),
     subimos el archivo a un bucket `reports/` con TTL implícito (no borramos,
     pero los nombres incluyen UUID + timestamp para no colisionar) y devolvemos
     una signed URL pública por 24h.
  2. Fallback: guardar en backend/data/reports/ y servir como /static/reports/<nombre>.
     (El fallback es para dev local; en prod se usa Supabase.)

El consumidor recibe siempre:
    {
      "ok": true,
      "url": "https://...",
      "filename": "resumen-proyecto-2026-05-10.pdf",
      "share_message": "Hola, te comparto el resumen del proyecto: <url>"
    }
"""
from __future__ import annotations

import logging
import uuid
from datetime import date as Date
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from models.payment import Payment
from models.milestone import Milestone
from models.project import Project
from models.user import User
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

REPORTS_DIR = Path(__file__).resolve().parent.parent / "data" / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


async def _gather_data(
    db: AsyncSession,
    user: User,
    scope: str,
    period_from: str | None,
    period_to: str | None,
) -> dict[str, Any]:
    """Reúne los datos necesarios según el scope."""
    data: dict[str, Any] = {"user_email": user.email}

    if scope in ("project", "full"):
        proj = (
            await db.execute(select(Project).where(Project.user_id == user.id))
        ).scalar_one_or_none()
        data["project"] = proj

    if scope in ("payments", "full"):
        q = select(Payment).where(Payment.user_id == user.id)
        if period_from:
            q = q.where(Payment.fecha >= Date.fromisoformat(period_from))
        if period_to:
            q = q.where(Payment.fecha <= Date.fromisoformat(period_to))
        q = q.order_by(Payment.fecha.desc()).limit(500)
        data["payments"] = list((await db.execute(q)).scalars().all())

    if scope in ("milestones", "full"):
        ms = list(
            (
                await db.execute(
                    select(Milestone)
                    .where(Milestone.user_id == user.id)
                    .order_by(Milestone.due_date.asc().nulls_last())
                    .limit(500)
                )
            )
            .scalars()
            .all()
        )
        data["milestones"] = ms

    return data


def _money(v: Any) -> str:
    if v is None:
        return "-"
    try:
        return f"$ {float(v):,.0f}".replace(",", ".")
    except (TypeError, ValueError):
        return str(v)


# ════════════════════════════════════════════════════════════════════
# PDF (reportlab)
# ════════════════════════════════════════════════════════════════════
def _render_pdf(title: str, scope: str, data: dict[str, Any]) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=title,
    )
    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h1.textColor = colors.HexColor("#0f172a")
    h2 = ParagraphStyle(
        "h2", parent=styles["Heading2"], textColor=colors.HexColor("#475569")
    )
    body = ParagraphStyle(
        "body", parent=styles["BodyText"], fontSize=10, leading=14
    )

    story = []
    story.append(Paragraph(title, h1))
    story.append(
        Paragraph(
            f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')} — RE Expert",
            body,
        )
    )
    story.append(Spacer(1, 0.6 * cm))

    if scope in ("project", "full") and data.get("project"):
        proj = data["project"]
        story.append(Paragraph("Estado del proyecto", h2))
        rows = [
            ["Nombre", proj.nombre or "-"],
            ["Estado", f"{proj.estado} — {proj.estado_texto}"],
            ["Presupuesto base", _money(proj.presupuesto_base)],
            ["Costo real", _money(proj.costo_real)],
            ["Avance real", f"{float(proj.avance_real_pct):.1f} %"],
            ["Avance plan", f"{float(proj.avance_plan_pct):.1f} %"],
            ["Plazo", f"{proj.meses_transcurridos}/{proj.meses_total} meses"],
        ]
        if proj.fecha_inicio:
            rows.append(["Inicio", proj.fecha_inicio.isoformat()])
        if proj.fecha_entrega_programada:
            rows.append(
                ["Entrega programada", proj.fecha_entrega_programada.isoformat()]
            )
        if proj.fecha_entrega_estimada:
            rows.append(
                ["Entrega estimada", proj.fecha_entrega_estimada.isoformat()]
            )
        t = Table(rows, colWidths=[5 * cm, 11 * cm])
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f1f5f9")),
                    ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#475569")),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e2e8f0")),
                ]
            )
        )
        story.append(t)
        story.append(Spacer(1, 0.6 * cm))

    if scope in ("payments", "full") and data.get("payments") is not None:
        story.append(Paragraph("Pagos", h2))
        rows: list[list[Any]] = [
            ["Fecha", "Concepto", "Proveedor", "Monto", "Estado"]
        ]
        total = 0.0
        for p in data["payments"]:
            rows.append(
                [
                    p.fecha.isoformat(),
                    p.concepto[:60],
                    (p.proveedor or "-")[:30],
                    _money(p.monto),
                    p.estado,
                ]
            )
            try:
                total += float(p.monto)
            except (TypeError, ValueError):
                pass
        rows.append(["", "", "TOTAL", _money(total), ""])
        t = Table(rows, colWidths=[2.2 * cm, 6.4 * cm, 4 * cm, 2.7 * cm, 2 * cm])
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("ALIGN", (3, 1), (3, -1), "RIGHT"),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e2e8f0")),
                    ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#f8fafc")),
                    ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
                ]
            )
        )
        story.append(t)
        story.append(Spacer(1, 0.6 * cm))

    if scope in ("milestones", "full") and data.get("milestones"):
        story.append(Paragraph("Cronograma", h2))
        rows = [["Hito", "Estado", "Fecha objetivo"]]
        for m in data["milestones"]:
            rows.append(
                [
                    m.name[:60],
                    m.status,
                    m.due_date.isoformat() if m.due_date else "-",
                ]
            )
        t = Table(rows, colWidths=[9 * cm, 3 * cm, 3.5 * cm])
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e2e8f0")),
                ]
            )
        )
        story.append(t)

    doc.build(story)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════════
# DOCX (python-docx)
# ════════════════════════════════════════════════════════════════════
def _render_docx(title: str, scope: str, data: dict[str, Any]) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)
    doc.add_paragraph(
        f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')} — RE Expert"
    )

    if scope in ("project", "full") and data.get("project"):
        proj = data["project"]
        doc.add_heading("Estado del proyecto", level=1)
        kv = [
            ("Nombre", proj.nombre or "-"),
            ("Estado", f"{proj.estado} — {proj.estado_texto}"),
            ("Presupuesto base", _money(proj.presupuesto_base)),
            ("Costo real", _money(proj.costo_real)),
            ("Avance real", f"{float(proj.avance_real_pct):.1f} %"),
            ("Avance plan", f"{float(proj.avance_plan_pct):.1f} %"),
            ("Plazo", f"{proj.meses_transcurridos}/{proj.meses_total} meses"),
        ]
        t = doc.add_table(rows=len(kv), cols=2)
        t.style = "Light Grid Accent 1"
        for i, (k, v) in enumerate(kv):
            t.cell(i, 0).text = k
            t.cell(i, 1).text = str(v)

    if scope in ("payments", "full") and data.get("payments") is not None:
        doc.add_heading("Pagos", level=1)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Fecha", "Concepto", "Proveedor", "Monto", "Estado"]):
            hdr[i].text = h
        total = 0.0
        for p in data["payments"]:
            r = t.add_row().cells
            r[0].text = p.fecha.isoformat()
            r[1].text = p.concepto
            r[2].text = p.proveedor or "-"
            r[3].text = _money(p.monto)
            r[4].text = p.estado
            try:
                total += float(p.monto)
            except (TypeError, ValueError):
                pass
        r = t.add_row().cells
        r[2].text = "TOTAL"
        r[3].text = _money(total)

    if scope in ("milestones", "full") and data.get("milestones"):
        doc.add_heading("Cronograma", level=1)
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Hito", "Estado", "Fecha objetivo"]):
            hdr[i].text = h
        for m in data["milestones"]:
            r = t.add_row().cells
            r[0].text = m.name
            r[1].text = m.status
            r[2].text = m.due_date.isoformat() if m.due_date else "-"

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════════
# Storage: Supabase si está configurado, sino disco local.
# ════════════════════════════════════════════════════════════════════
async def _upload_to_supabase(filename: str, blob: bytes, content_type: str) -> str | None:
    """Sube a `reports/` y devuelve URL pública firmada por 24h. None si falla."""
    try:
        import httpx

        from config.settings import settings

        if not (settings.SUPABASE_URL and settings.SUPABASE_SERVICE_ROLE_KEY):
            return None
        base = settings.SUPABASE_URL.rstrip("/")
        bucket = "reports"
        key = filename
        url_upload = f"{base}/storage/v1/object/{bucket}/{key}"
        headers = {
            "Authorization": f"Bearer {settings.SUPABASE_SERVICE_ROLE_KEY}",
            "apikey": settings.SUPABASE_SERVICE_ROLE_KEY,
            "Content-Type": content_type,
            "x-upsert": "true",
        }
        async with httpx.AsyncClient(timeout=30) as cli:
            r = await cli.post(url_upload, headers=headers, content=blob)
            if r.status_code >= 400:
                logger.warning("Supabase upload %s -> %s", r.status_code, r.text[:200])
                return None
            # Crear signed URL (24h)
            sign_url = f"{base}/storage/v1/object/sign/{bucket}/{key}"
            r2 = await cli.post(
                sign_url,
                headers={
                    "Authorization": f"Bearer {settings.SUPABASE_SERVICE_ROLE_KEY}",
                    "apikey": settings.SUPABASE_SERVICE_ROLE_KEY,
                    "Content-Type": "application/json",
                },
                json={"expiresIn": 24 * 3600},
            )
            if r2.status_code >= 400:
                logger.warning("Sign URL failed %s", r2.text[:200])
                return None
            payload = r2.json()
            signed_path = payload.get("signedURL") or payload.get("signedUrl")
            if not signed_path:
                return None
            if signed_path.startswith("/"):
                return f"{base}/storage/v1{signed_path}"
            return f"{base}/storage/v1/{signed_path}"
    except Exception as e:
        logger.exception("Supabase upload error: %s", e)
        return None


_LOCAL_HOSTS = ("localhost", "127.0.0.1", "0.0.0.0")


def _is_publicly_reachable(url: str) -> bool:
    """Un link sólo sirve para mandarle a OTRA persona si no apunta a localhost:
    los reports en disco local no son alcanzables desde el celular del destinatario."""
    u = (url or "").lower()
    return bool(u) and not any(h in u for h in _LOCAL_HOSTS)


def _save_local(filename: str, blob: bytes) -> str:
    """Guarda el PDF en disco y devuelve URL ABSOLUTA (necesario para que el
    link funcione cuando se manda por WhatsApp/Telegram al destinatario)."""
    from config.settings import settings

    p = REPORTS_DIR / filename
    p.write_bytes(blob)
    base = settings.BACKEND_PUBLIC_URL.rstrip("/") if settings.BACKEND_PUBLIC_URL else ""
    url = f"{base}/static/reports/{filename}"
    if not settings.DEBUG and not _is_publicly_reachable(url):
        logger.warning(
            "Report guardado local con URL no pública (%s): no será alcanzable "
            "para compartir. Configurá Supabase Storage o BACKEND_PUBLIC_URL público.",
            url,
        )
    return url


async def generate_report(
    db: AsyncSession,
    user: User,
    fmt: str = "pdf",
    scope: str = "project",
    title: str = "Resumen RE Expert",
    period_from: str | None = None,
    period_to: str | None = None,
) -> dict[str, Any]:
    """Genera el documento, lo sube/guarda y devuelve URL + mensaje compartible."""
    if fmt not in ("pdf", "docx"):
        return {"error": f"Formato no soportado: {fmt}"}
    if scope not in ("project", "payments", "milestones", "full"):
        return {"error": f"Scope no soportado: {scope}"}

    data = await _gather_data(db, user, scope, period_from, period_to)

    if fmt == "pdf":
        blob = _render_pdf(title, scope, data)
        ct = "application/pdf"
        ext = "pdf"
    else:
        blob = _render_docx(title, scope, data)
        ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"

    today = datetime.now().strftime("%Y-%m-%d")
    filename = f"{scope}-{today}-{uuid.uuid4().hex[:8]}.{ext}"

    url = await _upload_to_supabase(filename, blob, ct)
    from_supabase = url is not None
    if not url:
        url = _save_local(filename, blob)

    reachable = from_supabase or _is_publicly_reachable(url)
    share_msg = (
        f"Hola, te comparto el {title.lower()} generado desde RE Expert:\n{url}"
    )
    result: dict[str, Any] = {
        "ok": True,
        "url": url,
        "filename": filename,
        "format": fmt,
        "scope": scope,
        "share_message": share_msg,
        "publicly_reachable": reachable,
    }
    if not reachable:
        result["warning"] = (
            "El documento quedó en disco local (BACKEND_PUBLIC_URL apunta a "
            "localhost): este link NO le va a funcionar a quien se lo mandes. "
            "Para compartir, configurá Supabase Storage o un BACKEND_PUBLIC_URL público."
        )
    return result
